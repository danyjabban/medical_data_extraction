from docx import Document


def post_processor(answers:dict):
    """
    process answers so they can be loaded into document in the desired format
    ------
    Args:
        answers (dict): unprocessed dictionary of answers to each question
    Returns:
        None
    """
    processed_answers = process_answers(answers)
    create_doc(processed_answers)


def process_answers(answers:dict):
    """
    process answer into a stardard dictionary format
    ------
    Args:
        answers (dict): unprocessed dictionary of answers to each question
    Returns:
        processed_answers (dict): processed dictionary of answers to each question
    """
    processed_answers = {}
    for query in answers.keys():
        processed_answers[query] = {}
        for page_ans in answers[query]:
            page_idx = page_ans[0]
            for k, v in page_ans[1].items():
                k_standard = convert_str_standard(k)
                # check if key is already an answer, if so just add the page number
                if k_standard in processed_answers[query]:
                    cur = v
                    prev = processed_answers[query][k_standard]
                    processed_answers[query][k_standard][1].add(page_idx)
                else:
                    processed_answers[query][k_standard] = [v, {page_idx}]
    return processed_answers


def convert_str_standard(string:str):
    """
    convert each string to a stadard format
    """
    string = string.lower()
    string = string.replace('\n',' ') 
    string = string.strip()
    return string


def create_doc(answers:dict):
    """
    takes in a dictionary of answers and formats them into a word document
    """
    document = Document()
    for k, v in answers.items():
        document.add_heading(f'{k}', level=1)
        # lazy way of checking which query it is and then following that queries layout
        if 'medication' in k:
            for ans_k, ans_v in answers[k].items():
                document.add_paragraph(f'medication: {ans_k}\n start and stop dates: {ans_v[0]}\n page(s) found on: {ans_v[1]}')
        elif 'surgeries' in k:
            for ans_k, ans_v in answers[k].items():
                document.add_paragraph(f'surgeries: {ans_k}\n date: {ans_v[0]}\n page(s) found on: {ans_v[1]}')
        elif 'allergies' in k:
            for ans_k, ans_v in answers[k].items():
                document.add_paragraph(f'allergies: {ans_k}\n page(s) found on: {ans_v[1]}')
    document.save('docx_file.docx')
    

if __name__ == "__main__":
    # d = {'medication': [{51: {'Cymbalta': (None, None), 'temazepam': (None, None), 'trazodone': (None, None)}}, {111: {'Citalopram (Celexa) 20 MG Tab': ('04/10/2018', '10/20/2018'), 'Norethindrone Ac-Eth Estradiol ( Norethind-Eth Estrad 1-0.02 Mg) 1 Each Tablet': ('04/04/2018', '04/10/2018'), 'Trazodone (Desyrel) 50 MG Tab': ('12/28/2017', '07/16/2018'), 'Citalopram (Celexa) 10 MG Tab': ('12/28/2017', '04/10/2018'), 'Trazodone': ('12/28/2017', '07/16/2018'), 'Citalopram': ('10/19/2017', '12/28/2017'), 'Hydroxyzine Hcl': ('12/06/2017', '12/28/2017'), 'Trazodone Hcl': ('12/06/2017', '12/28/2017')}}]}
    d = {'medication': [[51, {'Cymbalta': (None, None), 'temazepam': (None, None), 'trazodone': (None, None)}], [111, {'Citalopram (Celexa) 20 MG Tab': [('07/16/2018', '09/18/2018'), ('04/20/2018', '07/16/2018'), ('04/20/2018', '04/20/2018'), ('04/10/2018', '10/4/20/2018')], 'Citalopram (Celexa) 20 MG Tab\nNorethindrone Ac-Eth Estradiol\n(Norethind-Eth Estrad 1-0.02 Mg) 1 EA\nTablet': [('04/04/2018', '04/10/2018')], 'Norethindrone Ac-Eth Estradiol\n(Norethind-Eth Estrad 1-0.02 Mg) 1 EA\nTablet': [('04/04/2018', '04/10/2018')], 'Trazodone (Desyrel) 50 MG Tab': [('12/28/2017', '07/16/2018')], 'Citalopram (Celexa) 10 MG Tab': [('12/28/2017', '07/16/2018'), ('12/28/2017', '04/10/2018')], 'Trazodone': [('12/28/2017', '07/16/2018')], 'Citalopram': [('12/28/2017', '04/10/2018'), ('10/19/2017', '12/28/2017')], 'Hydroxyzine Hcl': [('12/06/2017', '12/28/2017'), ('12/06/2017', '12/28/2017'), ('12/06/2017', '12/28/2017'), ('12/06/2017', '12/28/2017'), ('12/06/2017', '12/28/2017')], 'Trazodone Hcl': [('12/06/2017', '12/28/2017'), ('12/06/2017', '12/28/2017'), ('12/06/2017', '12/28/2017')]}]]}
    # process_answers(d)
    post_processor(d)